# utils/docx_generator.py
import io
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from loguru import logger

from core.profile_names import has_arabic
from utils.cv_context import build_cv_context
from utils.cv_photo import data_uri_to_bytes
from utils.docx_styles import resolve_docx_style
from utils.template_registry import DEFAULT_TEMPLATE_ID

OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Word renders Arabic shaping/bidi natively, so unlike the PDF generator we
# don't need reshaping here — just a complex-script font and the RTL
# paragraph/run flags below.
ARABIC_FONT = "Traditional Arabic"


def _set_rtl_paragraph(paragraph, is_arabic: bool):
    if not is_arabic:
        return
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _set_rtl_run(run, is_arabic: bool):
    if not is_arabic:
        return
    run.font.name = ARABIC_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), ARABIC_FONT)
    rtl = rPr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        rPr.append(rtl)


def _add_run(paragraph, text, is_arabic, style, size=10.5, bold=False, color=None):
    r = paragraph.add_run(text)
    r.font.size = Pt(size)
    r.font.name = style["font_name"] if not is_arabic else ARABIC_FONT
    r.bold = bold
    if color:
        r.font.color.rgb = color
    _set_rtl_run(r, is_arabic)
    return r


def _add_bottom_border(paragraph, color_hex: str):
    """Adds a single-rule bottom border to a paragraph, in the given accent
    color — gives a heading real visual presence in a format with no CSS.
    See docx_styles.py's module docstring for why this exists."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")     # eighths of a point
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _shade_paragraph(paragraph, color_hex: str):
    """Fills a paragraph's background — used for templates whose PDF
    identity IS a solid color block behind the name (sidebar_dark,
    bold_banner). See docx_styles.py's header_shade docstring."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#"))
    pPr.append(shd)


def _heading(doc, title, is_arabic, style):
    p = doc.add_heading(title, level=2)
    run = p.runs[0]
    run.font.name = style["font_name"] if not is_arabic else ARABIC_FONT
    run.font.color.rgb = style["heading_color_rgb"]
    run.underline = style["heading_underline"]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    if style.get("heading_border"):
        _add_bottom_border(p, style["accent_color"])
    _set_rtl_paragraph(p, is_arabic)
    _set_rtl_run(run, is_arabic)


def _bullet(doc, text, is_arabic, style):
    p = doc.add_paragraph(text, style=style["bullet_style"])
    p.runs[0].font.size = Pt(10.5)
    p.runs[0].font.name = style["font_name"] if not is_arabic else ARABIC_FONT
    p.paragraph_format.space_after = Pt(2)
    _set_rtl_paragraph(p, is_arabic)
    _set_rtl_run(p.runs[0], is_arabic)


def _verbatim_is_rtl(text: str, is_arabic: bool) -> bool:
    """
    Should this verbatim line be laid out right-to-left?

    Only for the two fields no agent translates — publications and the CV's
    own sections. Word applies the paragraph's bidi flag to the neutral
    characters in a line, so an untranslated English line inside an Arabic
    document comes out with its leading number pushed to the far end, exactly
    as it did in the PDF (see the [dir="ltr"] rules in pdf_generator.py and
    utils/cv_context.text_direction). Flagging the line by its own script
    keeps the .docx and the .pdf showing the same thing.
    """
    return is_arabic and has_arabic(text)


def _add_photo(doc, photo: str | None, style: dict) -> None:
    """
    Puts the candidate's picture, centred, above the name block.

    No-ops unless the preset declares `photo` AND there is one on file, so a
    photo template chosen by someone with no photo produces a normal header
    rather than a gap — the same graceful-degradation rule the HTML
    templates follow with `{% if photo %}`.

    Never raises: a picture that python-docx rejects costs the user the
    picture, not the .docx. The PDF beside it is rendered independently
    (main.py runs the three renders in parallel) and is unaffected either way.
    """
    if not photo or not style.get("photo"):
        return
    raw = data_uri_to_bytes(photo)
    if not raw:
        return
    try:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 1.25" tall is roughly the printed size the HTML templates use, and
        # height-constrained (not width) because the stored image keeps the
        # candidate's own aspect ratio — see utils/cv_photo.py, which scales
        # to fit a box rather than cropping to a fixed shape.
        paragraph.add_run().add_picture(io.BytesIO(raw), height=Inches(1.25))
        paragraph.paragraph_format.space_after = Pt(6)
    except Exception as e:
        logger.warning(f"📷 Couldn't embed the candidate photo in the .docx, continuing without it: {e}")


def generate_cv_docx(state: dict, output_path: str, template_id: str | None = None) -> str:
    """
    Generates .docx version of the tailored CV, styled to match the chosen
    template_id (see utils/docx_styles.py for what "match" means for a
    format that can't render arbitrary CSS). Falls back to the default
    style if template_id is None or unrecognized.

    output_path is REQUIRED (no static-filename fallback) — same reasoning
    as pdf_generator.py's render_cv_pdf/render_cover_letter_pdf: a fixed
    default path is exactly the bug that let concurrent users' generations
    overwrite each other. Every caller, including tests, must supply a
    unique per-request/per-call path.
    """
    doc = Document()
    resolved_template_id = template_id or DEFAULT_TEMPLATE_ID
    style = resolve_docx_style(resolved_template_id)

    context = build_cv_context(state, template_id=resolved_template_id)
    is_arabic = context["is_arabic"]
    personal = context["personal"]

    # Candidate photo, for the presets that declare one. context["photo"] is
    # already gated on the template having a slot (see cv_context's
    # resolve_candidate_photo), so this is just "is there one to draw".
    _add_photo(doc, context.get("photo"), style)

    # Header — templates with a header_shade get a filled color block behind
    # the name/contact block (evokes their PDF sidebar/banner identity);
    # header_text_color keeps the text readable against that fill.
    header_shade = style.get("header_shade")
    header_text_rgb = style.get("header_text_color_rgb") or style["heading_color_rgb"]

    name_p = doc.add_heading(personal.get("name", ""), level=1)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_rtl_paragraph(name_p, is_arabic)
    if name_p.runs:
        name_p.runs[0].font.name = style["font_name"] if not is_arabic else ARABIC_FONT
        name_p.runs[0].font.color.rgb = header_text_rgb if header_shade else style["heading_color_rgb"]
        _set_rtl_run(name_p.runs[0], is_arabic)
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # keep centered even after RTL flag
    if header_shade:
        _shade_paragraph(name_p, header_shade)

    # personal.linkedin / .github are bare handles (see _profile_handle in
    # utils/cv_context.py). The linkedin line was missing the "in/" path
    # segment the HTML templates all use, so even a correct handle printed a
    # URL that doesn't resolve — the .docx and the .pdf of the same CV showed
    # two different addresses.
    contact = " · ".join(filter(None, [
        personal.get("email", ""),
        personal.get("location", ""),
        f"linkedin.com/in/{personal['linkedin']}" if personal.get("linkedin") else "",
        f"github.com/{personal['github']}" if personal.get("github") else "",
    ]))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(cp, contact, is_arabic, style, size=9.5, color=header_text_rgb if header_shade else None)
    _set_rtl_paragraph(cp, is_arabic)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header_shade:
        _shade_paragraph(cp, header_shade)
        cp.paragraph_format.space_after = Pt(10)

    # Section headings for everything added with the facts_schema expansion
    # come from context["labels"] (utils/cv_context.py) rather than inline
    # literals, so the .docx and the .pdf of the same CV can't drift apart on
    # what a section is called. The older headings below are left as they are.
    labels = context.get("labels", {})

    # Summary
    if context["tailored_summary"]:
        _heading(doc, "الملخص المهني" if is_arabic else "Professional Summary", is_arabic, style)
        p = doc.add_paragraph()
        _add_run(p, context["tailored_summary"], is_arabic, style)
        _set_rtl_paragraph(p, is_arabic)

    # Key achievements sit directly under the summary — the career-highlights
    # position, and the same order the HTML templates use, so the two formats
    # of one CV read identically.
    if context.get("major_achievements"):
        _heading(doc, labels.get("achievements", "Key Achievements"), is_arabic, style)
        for item in context["major_achievements"]:
            _bullet(doc, item, is_arabic, style)

    # Experience — context["experience"][i]["bullets"] is already the
    # resolved (tailored, in original order) list of strings.
    if context["experience"]:
        _heading(doc, "الخبرة العملية" if is_arabic else "Experience", is_arabic, style)
        for exp in context["experience"]:
            p = doc.add_paragraph()
            _add_run(p, f"{exp['title']}   {exp['company']}", is_arabic, style, bold=True)
            _add_run(p, f"    {exp['dates']}", is_arabic, style, size=10, color=style["accent_color_rgb"])
            _set_rtl_paragraph(p, is_arabic)
            for bullet_text in exp["bullets"]:
                _bullet(doc, bullet_text, is_arabic, style)

    # Projects — one tailored description paragraph per project, matching
    # the real tailoring_engine.py output shape (not a bullet list).
    if context["projects"]:
        _heading(doc, "المشاريع" if is_arabic else "Projects", is_arabic, style)
        for proj in context["projects"]:
            p = doc.add_paragraph()
            _add_run(p, proj["name"], is_arabic, style, bold=True)
            tech_stack = ", ".join(proj["tech_stack"])
            if tech_stack:
                _add_run(p, f"    {tech_stack}", is_arabic, style, size=10, color=style["accent_color_rgb"])
            _set_rtl_paragraph(p, is_arabic)
            if proj["description"]:
                dp = doc.add_paragraph()
                _add_run(dp, proj["description"], is_arabic, style, size=10.5)
                _set_rtl_paragraph(dp, is_arabic)

    # Skills
    if context["skills"]:
        _heading(doc, "المهارات" if is_arabic else "Skills", is_arabic, style)
        # Sourced from context["labels"] (utils/cv_context.py) rather than a
        # local dict, so the DOCX and the PDF of the same CV can't disagree
        # about what a section is called — they previously had two separate
        # Arabic wordings for the same skill categories.
        label_keys = {
            "languages": "languages", "frameworks": "frameworks", "tools": "tools",
            "soft_skills": "soft_skills", "other": "other_skills",
        }
        for category, items in context["skills"].items():
            if items:
                label = (
                    labels.get(label_keys.get(category, ""), category)
                    if is_arabic
                    else category.replace("_", " ").capitalize()
                )
                p = doc.add_paragraph()
                _add_run(p, f"{label}: ", is_arabic, style, bold=True)
                _add_run(p, ", ".join(items), is_arabic, style)
                p.paragraph_format.space_after = Pt(2)
                _set_rtl_paragraph(p, is_arabic)

    # Human languages, straight after Skills — where the HTML templates put
    # them, and where a reader expects a language list.
    if context.get("languages_spoken"):
        _heading(doc, labels.get("spoken_languages", "Languages"), is_arabic, style)
        p = doc.add_paragraph()
        _add_run(p, ", ".join(context["languages_spoken"]), is_arabic, style)
        _set_rtl_paragraph(p, is_arabic)

    # Volunteer work
    if context["volunteer_work"]:
        _heading(doc, "الأعمال التطوعية" if is_arabic else "Volunteer Work", is_arabic, style)
        for v in context["volunteer_work"]:
            _bullet(doc, v, is_arabic, style)

    # Education
    if context["education"]:
        _heading(doc, "التعليم" if is_arabic else "Education", is_arabic, style)
        for edu in context["education"]:
            p = doc.add_paragraph()
            _add_run(p, f"{edu.get('degree', '')}   {edu.get('institution', '')}", is_arabic, style, bold=True)
            _add_run(p, f"    {edu.get('graduation_year', '')}", is_arabic, style, size=10)
            _set_rtl_paragraph(p, is_arabic)
            if edu.get("gpa"):
                gp = doc.add_paragraph()
                _add_run(gp, f"{'المعدل التراكمي' if is_arabic else 'GPA'}: {edu.get('gpa', '')}", is_arabic, style, size=10, color=style["accent_color_rgb"])
                gp.paragraph_format.space_after = Pt(2)
                _set_rtl_paragraph(gp, is_arabic)

    # Training and courses — a course attended, as opposed to a credential
    # held (that stays under Certifications, immediately below).
    if context.get("training_courses"):
        _heading(doc, labels.get("training", "Training & Courses"), is_arabic, style)
        for course in context["training_courses"]:
            _bullet(doc, ", ".join(filter(None, [
                course.get("name"), course.get("provider"), course.get("date"),
            ])), is_arabic, style)

    # Certifications
    if context["certifications"]:
        _heading(doc, "الشهادات" if is_arabic else "Certifications", is_arabic, style)
        for cert in context["certifications"]:
            _bullet(doc, cert, is_arabic, style)

    if context.get("publications"):
        _heading(doc, labels.get("publications", "Publications"), is_arabic, style)
        for pub in context["publications"]:
            citation = ", ".join(filter(None, [
                pub.get("title"), pub.get("venue"), pub.get("year"),
            ]))
            _bullet(doc, citation, _verbatim_is_rtl(citation, is_arabic), style)

    if context.get("participation"):
        _heading(doc, labels.get("participation", "Conferences & Participation"), is_arabic, style)
        for item in context["participation"]:
            _bullet(doc, ", ".join(filter(None, [
                item.get("title"), item.get("role"), item.get("organization"),
                item.get("scope"), item.get("date"),
            ])), is_arabic, style)

    if context.get("teaching_and_editorial"):
        _heading(doc, labels.get("teaching_editorial", "Teaching & Editorial Boards"), is_arabic, style)
        for item in context["teaching_and_editorial"]:
            _bullet(doc, item, is_arabic, style)

    if context.get("awards"):
        _heading(doc, labels.get("awards", "Awards"), is_arabic, style)
        for award in context["awards"]:
            _bullet(doc, award, is_arabic, style)

    # The CV's own sections, under the candidate's own headings — see
    # FactsJSON.additional_sections. Verbatim: this is where a procedure
    # count or a flight-hours total lives, and nothing here reformats one.
    for section in context.get("additional_sections", []):
        # cv_context already decided this section's reading direction from its
        # own content — reuse it rather than re-deriving a second opinion.
        section_rtl = is_arabic and section.get("dir") == "rtl"
        _heading(doc, section["section_title"], section_rtl, style)
        for entry in section["entries"]:
            _bullet(doc, entry, section_rtl, style)

    doc.save(output_path)
    logger.info(f"✅ CV DOCX saved → {output_path} (template: {resolved_template_id})")
    return output_path
