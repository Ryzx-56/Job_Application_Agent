"""
The FactsJSON expansion, end to end through both renderers.

WHAT THIS GUARDS. FactsJSON gained named fields for content real CVs carry
(achievements, courses, participation, publications, teaching/editorial,
awards, languages) plus `additional_sections`, the catch-all that keeps a CV
section nothing else covers under the candidate's own heading. Several
guarantees around that are invisible in isolation and only fail at render
time, which is why these tests drive the real templates:

  · nothing is dropped — every field reaches every one of the 15 templates
  · nothing is rewritten — the catch-all and publications render verbatim,
    including their numbers, in Arabic as well as English
  · a template that carries content it doesn't print trips fit_to_page's
    clipping guard, which is FATAL (refunds the credit), so section coverage
    is a correctness test and not a cosmetic one
  · one canonical section order across all templates

Deliberately NO LLM CALLS: everything below runs off a fixed facts_json, so
it is fast, free and deterministic. The one test that does exercise Agent 1's
routing is marked `extra_api_call` and is excluded by default (see pytest.ini).
"""
import fitz
import pytest
from docx import Document

import main as api
from agents.cv_parser import parse_cv_text
from schemas.facts_schema import FactsJSON
from utils.cv_context import _profile_handle, build_cv_context, text_direction
from utils.docx_generator import generate_cv_docx
from utils.fit_to_page import (CLIPPING_MIN_RATIO, _context_text_length,
                               _glyph_count)
from utils.pdf_generator import render_cv_pdf
from utils.template_registry import TEMPLATE_REGISTRY

# A CV with content in every new field, two sections that map to NO named
# field, and contact details in the shape people actually paste them (full
# URLs, which is what produced linkedin.com/in/https://...).
FACTS = {
    "personal": {
        "name": "Yousef Al-Harbi",
        "email": "yousef@example.com",
        "phone": "+966 55-123-4567",
        "linkedin": "https://www.linkedin.com/in/yousef-alharbi/",
        "github": "https://github.com/yalharbi",
        "location": "Riyadh, Saudi Arabia",
    },
    "summary": "Consultant cardiothoracic surgeon with 11 years of operative experience.",
    "education": [{"institution": "King Saud University", "degree": "MBBS",
                   "gpa": "4.6", "graduation_year": "2011", "distinctions": [],
                   "relevant_coursework": []}],
    "experience": [{"company": "King Faisal Specialist Hospital",
                    "title": "Consultant Surgeon", "dates": "2018 - Present",
                    "bullets": ["Led the adult cardiac surgery service"], "metrics": []}],
    "skills": {"languages": [], "frameworks": [], "tools": ["Clinical audit"],
               "soft_skills": [], "other": ["Valve repair"]},
    "projects": [],
    "certifications": ["Saudi Board of Cardiothoracic Surgery"],
    "languages_spoken": ["Arabic (native)", "English (fluent)"],
    "volunteer_work": [],
    "awards": ["Ministry of Health Excellence Award, 2022"],
    "major_achievements": ["Established the first minimally invasive valve programme"],
    "training_courses": [{"name": "Advanced Trauma Life Support",
                          "provider": "American College of Surgeons", "date": "2023"}],
    "participation": [{"title": "Saudi Heart Association Conference", "role": "Speaker",
                       "organization": "Saudi Heart Association", "scope": "local",
                       "date": "2024"}],
    "publications": [{"title": "Outcomes of Minimally Invasive Mitral Valve Repair",
                      "venue": "Saudi Medical Journal", "year": "2024", "url": None}],
    "teaching_and_editorial": ["Clinical lecturer, King Saud University, 2019-Present"],
    "additional_sections": [
        {"section_title": "Surgical Outcomes",
         "entries": ["1,240 procedures performed as primary operator between 2019 and 2024.",
                     "0.8% 30-day complication rate across all cardiothoracic cases."]},
        {"section_title": "Flight Operations Record",
         "entries": ["4,800 total flight hours, of which 1,150 as pilot in command."]},
    ],
}

# Every figure that must survive from facts_json to the rendered page byte for
# byte. A paraphrased number is a fabricated number.
FIGURES = ["1,240", "0.8%", "4,800", "1,150"]

# One line of proof per section, chosen so a heading rendered without its
# content (or content without its heading) still fails.
SECTION_EVIDENCE = {
    "achievements": "minimally invasive valve programme",
    "training": "Advanced Trauma Life Support",
    "certifications": "Saudi Board of Cardiothoracic Surgery",
    "publications": "Outcomes of Minimally Invasive Mitral Valve Repair",
    "participation": "Saudi Heart Association Conference",
    "teaching": "Clinical lecturer",
    "awards": "Ministry of Health Excellence Award",
    "languages": "Arabic (native)",
    # Body text, not the headings: two templates letter-space their section
    # titles, which extracts as "S U R G I C A L  O U T C O M E S". The
    # headings themselves are covered by the order assertion below, which
    # squashes whitespace before matching.
    "catch-all 1": "procedures performed as primary operator",
    "catch-all 2": "total flight hours",
}

# The canonical order every template renders in. Two-column templates emit
# their side column first, so only the pairs below are asserted rather than a
# single fixed sequence — these are the relationships that carry meaning.
_HEADINGS = {
    "professionalsummary": "SUM", "profile": "SUM", "summary": "SUM",
    "keyachievements": "ACH", "experience": "EXP", "workexperience": "EXP",
    "professionalexperience": "EXP", "projects": "PRJ", "skills": "SKL",
    "languages": "LNG", "education": "EDU", "training&courses": "TRN",
    "certifications": "CRT", "publications": "PUB",
    "conferences&participation": "PAR", "teaching&editorialboards": "TCH",
    "awards": "AWD", "surgicaloutcomes": "CAT1", "flightoperationsrecord": "CAT2",
}
ORDER_PAIRS = [("SUM", "ACH"), ("ACH", "EXP"), ("SKL", "LNG"),
               ("AWD", "CAT1"), ("CAT1", "CAT2")]
# Certifications sit in the side column of these two, i.e. before everything
# in the main column — so "training then certifications" cannot apply.
TWO_COLUMN = {"sidebar_dark", "portrait_rail"}


def _state(**overrides) -> dict:
    state = {"facts_json": FACTS, "cv_language": "en", "template_id": None,
             "profile_name_en": FACTS["personal"]["name"], "profile_name_ar": ""}
    state.update(overrides)
    return state


def _pdf_lines(path) -> list[str]:
    with fitz.open(path) as doc:
        return [line for page in doc for line in page.get_text().split("\n")]


def _pdf_text(path) -> str:
    """Whitespace-collapsed page text.

    Collapsed because a narrow column wraps a citation mid-phrase, and a
    letter-spaced heading extracts as "K E Y  A C H I E V E M E N T S" — both
    would read as missing content otherwise.
    """
    return " ".join(" ".join(_pdf_lines(path)).split())


def _heading_sequence(lines) -> list[str]:
    sequence = []
    for line in lines:
        key = _HEADINGS.get("".join(line.split()).lower())
        if key and (not sequence or sequence[-1] != key):
            sequence.append(key)
    return sequence


# ─── SCHEMA ──────────────────────────────────────────────────────────────────

def test_facts_json_round_trips():
    assert FactsJSON.model_validate(FACTS).model_dump()["additional_sections"][0][
        "section_title"] == "Surgical Outcomes"


def test_a_cv_saved_before_the_expansion_still_validates():
    """Stored generation_snapshots predate every field above and are replayed
    on any re-render, so the schema has to accept a facts_json without them."""
    legacy = FactsJSON.model_validate({"personal": {"name": "Old Snapshot"}})
    assert legacy.additional_sections == [] and legacy.languages_spoken == []


def test_odd_shapes_are_flattened_rather_than_rejected():
    """An LLM occasionally returns an object where a string was specified.
    Failing validation would burn Agent 1's retries and cost the user the whole
    CV over one badly-shaped list item, so those items are flattened."""
    facts = FactsJSON.model_validate({
        "personal": {"name": "T"},
        "awards": [{"name": "Best Paper", "year": "2024"}],
        "additional_sections": [{"entries": [{"metric": "1,240 procedures"}, "plain line"]}],
    })
    assert facts.awards == ["name: Best Paper, year: 2024"]
    assert facts.additional_sections[0].entries == ["metric: 1,240 procedures", "plain line"]
    # A heading the model failed to echo back is a formatting miss, not a
    # reason to throw away the content under it.
    assert facts.additional_sections[0].section_title == ""


# ─── CONTEXT ─────────────────────────────────────────────────────────────────

def test_context_carries_every_new_field():
    context = build_cv_context(_state())
    for key in ("major_achievements", "training_courses", "participation",
                "publications", "teaching_and_editorial", "awards",
                "languages_spoken", "additional_sections"):
        assert context[key], f"{key} missing from the render context"


def test_catch_all_and_publications_are_never_localized():
    """The two verbatim fields. A glossary substitution turned a '30-day
    complication rate' into a 'daily' one, and half-transliterated a citation
    into something nobody can look up — so neither field is translated at all."""
    glossary = {"Surgical Outcomes": "النتائج الجراحية",
                "day complication rate": "معدل المضاعفات اليومي",
                "Saudi Medical Journal": "المجلة الطبية السعودية",
                "Ministry of Health Excellence Award": "جائزة التميز من وزارة الصحة",
                "Arabic": "العربية"}
    context = build_cv_context(_state(cv_language="ar", arabic_glossary=glossary,
                                      profile_name_ar="يوسف الحربي"))

    catch_all = " ".join(e for s in context["additional_sections"] for e in s["entries"])
    assert "30-day complication rate" in catch_all
    assert context["additional_sections"][0]["section_title"] == "Surgical Outcomes"
    assert context["publications"][0]["venue"] == "Saudi Medical Journal"
    for figure in FIGURES:
        assert figure in catch_all or figure in " ".join(
            p["title"] for p in context["publications"])

    # Everything else still localizes — the exemption is narrow on purpose.
    assert "جائزة التميز" in " ".join(context["awards"])
    assert "العربية" in " ".join(context["languages_spoken"])


def test_verbatim_blocks_are_marked_with_their_own_direction():
    """Latin text inside an RTL page has its leading digits moved to the end of
    the line by the bidi algorithm — '1,240 procedures ...' renders as
    'procedures ... 1,240'. Each block carries the direction it reads in."""
    context = build_cv_context(_state(cv_language="ar", profile_name_ar="يوسف الحربي"))
    assert context["additional_sections"][0]["dir"] == "ltr"
    assert context["publications_dir"] == "ltr"

    assert text_direction(["قسم بالعربية"], "ltr") == "rtl"
    assert text_direction(["Surgical Outcomes"], "rtl") == "ltr"
    assert text_direction(["1,240"], "rtl") == "rtl"   # no strong character either way


def test_a_citation_does_not_print_its_own_full_stop_before_the_venue():
    """A title extracted from a full citation ends in the citation's own full
    stop, and the renderers then append ", {venue}" — printing
    "...Valve Repair., Saudi Medical Journal"."""
    facts = {**FACTS, "publications": [
        {"title": "Outcomes of Minimally Invasive Mitral Valve Repair.",
         "venue": "Saudi Medical Journal", "year": "2024", "url": None}]}
    publication = build_cv_context(_state(facts_json=facts))["publications"][0]
    assert publication["title"].endswith("Repair")
    # Nothing follows it -> the author's own punctuation is left alone.
    facts["publications"] = [{"title": "A Standalone Title.", "venue": "", "year": "", "url": None}]
    assert build_cv_context(_state(facts_json=facts))["publications"][0]["title"] == "A Standalone Title."


def test_blank_entries_never_reach_the_renderers():
    """docx_generator's _bullet() reads runs[0] on the paragraph it just made,
    and an empty string produces a paragraph with no runs — an IndexError and
    no .docx at all."""
    facts = {**FACTS, "awards": ["", "  "], "major_achievements": [""],
             "additional_sections": [{"section_title": "Empty", "entries": ["", " "]}]}
    context = build_cv_context(_state(facts_json=facts))
    assert context["awards"] == [] and context["major_achievements"] == []
    assert context["additional_sections"] == []


@pytest.mark.parametrize("raw,expected", [
    ("https://www.linkedin.com/in/yousef-alharbi/", "yousef-alharbi"),
    ("linkedin.com/in/yousef-alharbi", "yousef-alharbi"),
    ("in/yousef-alharbi", "yousef-alharbi"),
    ("@yousef-alharbi", "yousef-alharbi"),
    ("yousef-alharbi", "yousef-alharbi"),
    ("https://github.com/yalharbi", "yalharbi"),
    ("github.com/yalharbi/", "yalharbi"),
    ("https://www.linkedin.com/in/yousef-alharbi/?originalSubdomain=sa", "yousef-alharbi"),
    ("https://www.linkedin.com/", ""),
    ("", ""),
    (None, ""),
])
def test_profile_handle_normalizes_whatever_was_pasted(raw, expected):
    """Every template renders this behind a 'linkedin.com/in/' prefix of its
    own, so anything but a bare handle prints a URL nested inside a URL."""
    assert _profile_handle(raw) == expected


# ─── THE USABILITY GATE (main.py) ────────────────────────────────────────────

def test_gate_counts_the_new_fields_as_extracted_content():
    """The gate refuses a run whose facts_json holds too little of the source.
    A field it doesn't count is content that can never reach the numerator, so
    a CV whose substance is publications and its own sections would look
    emptier to the gate the better the extraction got."""
    without = {k: v for k, v in FACTS.items()
               if k not in ("publications", "additional_sections", "training_courses",
                            "participation", "teaching_and_editorial", "major_achievements",
                            "summary")}
    assert api._facts_content_length(FACTS) > api._facts_content_length(without)

    populated = api._populated_fact_groups(FACTS)
    for group in ("additional_sections", "publications", "training_courses",
                  "participation", "teaching_and_editorial", "major_achievements",
                  "languages_spoken", "summary"):
        assert group in populated, f"{group} is not counted as a populated fact group"


def test_gate_accepts_a_cv_whose_substance_is_the_new_fields():
    source = "x" * 4000  # long enough that the coverage ratio is consulted
    assert api._pipeline_produced_usable_cv(
        {"facts_json": FACTS, "raw_cv_text": source, "request_id": "test"})


# ─── RENDERING ───────────────────────────────────────────────────────────────

@pytest.mark.parametrize("template_id", list(TEMPLATE_REGISTRY))
def test_every_template_prints_every_section_in_the_canonical_order(template_id, tmp_path):
    path = str(tmp_path / f"{template_id}.pdf")
    render_cv_pdf(_state(), path, template_id=template_id)
    text, lines = _pdf_text(path), _pdf_lines(path)

    missing = [name for name, needle in SECTION_EVIDENCE.items() if needle.lower() not in text.lower()]
    assert not missing, f"{template_id} dropped: {missing}"
    for figure in FIGURES:
        assert figure in text, f"{template_id} lost the figure {figure}"

    sequence = _heading_sequence(lines)
    pairs = list(ORDER_PAIRS) + ([] if template_id in TWO_COLUMN else [("TRN", "CRT")])
    for first, second in pairs:
        assert first in sequence and second in sequence, f"{template_id}: {sequence}"
        assert sequence.index(first) < sequence.index(second), \
            f"{template_id} renders {second} before {first}: {sequence}"


@pytest.mark.parametrize("template_id", list(TEMPLATE_REGISTRY))
def test_every_template_clears_the_clipping_floor(template_id, tmp_path):
    """fit_to_page RAISES below this ratio, which refunds the credit and hands
    the user a failure — so a template carrying content it never prints is a
    production outage waiting for a long enough CV, not a cosmetic gap. Before
    the new sections were rolled out to all templates this sat at 57%."""
    path = str(tmp_path / f"{template_id}.pdf")
    render_cv_pdf(_state(), path, template_id=template_id)
    with open(path, "rb") as handle:
        ratio = _glyph_count(handle.read()) / _context_text_length(
            build_cv_context(_state(), template_id=template_id))
    assert ratio >= CLIPPING_MIN_RATIO, f"{template_id} drew only {ratio:.0%} of its content"


@pytest.mark.parametrize("template_id", list(TEMPLATE_REGISTRY))
def test_contact_urls_are_never_double_prefixed(template_id, tmp_path):
    """personal.linkedin arrives as a full URL and every template prefixes
    'linkedin.com/in/' onto it, which printed the address twice, nested."""
    path = str(tmp_path / f"{template_id}.pdf")
    render_cv_pdf(_state(), path, template_id=template_id)
    # "/ " closed up: a long contact line legitimately wraps AT the slash.
    text = _pdf_text(path).replace("/ ", "/")
    assert "linkedin.com/in/https" not in text and "github.com/https" not in text
    assert "linkedin.com/in/yousef-alharbi" in text
    with fitz.open(path) as doc:
        links = [link.get("uri", "") for page in doc for link in page.get_links()]
    assert not [link for link in links if link.count("http") > 1], links


def test_docx_prints_every_section_in_the_canonical_order(tmp_path):
    path = str(tmp_path / "cv.docx")
    generate_cv_docx(_state(), path)
    paragraphs = [p.text for p in Document(path).paragraphs]
    text = " ".join(paragraphs)

    for name, needle in SECTION_EVIDENCE.items():
        assert needle.lower() in text.lower(), f"DOCX dropped {name}"
    for figure in FIGURES:
        assert figure in text
    assert "linkedin.com/in/yousef-alharbi" in text and "github.com/yalharbi" in text

    sequence = _heading_sequence(paragraphs)
    for first, second in ORDER_PAIRS + [("TRN", "CRT")]:
        assert sequence.index(first) < sequence.index(second), sequence


def test_arabic_docx_lays_the_verbatim_blocks_out_left_to_right(tmp_path):
    """Same bidi problem as the PDF: Word applies the paragraph's direction to
    a line's neutral characters, so an untranslated English line in an Arabic
    document comes out with its leading number at the far end."""
    path = str(tmp_path / "cv_ar.docx")
    generate_cv_docx(_state(cv_language="ar", profile_name_ar="يوسف الحربي"), path)
    rtl_flag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi"

    checked = 0
    for paragraph in Document(path).paragraphs:
        if any(figure in paragraph.text for figure in FIGURES) or \
                paragraph.text.strip() in ("Surgical Outcomes", "Flight Operations Record"):
            assert paragraph._p.find(f".//{rtl_flag}") is None, \
                f"verbatim line laid out RTL: {paragraph.text[:60]}"
            checked += 1
    assert checked, "no verbatim lines found in the Arabic .docx"


# ─── AGENT 1'S ROUTING (opt in: pytest -m extra_api_call) ────────────────────

EXTRA_SECTIONS = """
MAJOR ACHIEVEMENTS:
- Ranked 1st out of 240 students in the college hackathon.

TRAINING AND COURSES:
- Deep Learning Specialization, DeepLearning.AI, 2024

PUBLICATIONS:
- Hawsawi, A. (2025). Multi-Agent Pipelines. Journal of Applied AI, 12(3).

SURGICAL OUTCOMES:
- 1,240 procedures performed as primary operator between 2019 and 2024.
- 0.8% 30-day complication rate across all cardiothoracic cases.
"""


@pytest.mark.extra_api_call
def test_agent1_routes_named_content_and_keeps_the_rest_under_its_own_heading():
    """The prompt's core instruction: use a named field when the content
    clearly matches, otherwise keep the section under the CV's own heading —
    and never copy the same content into both."""
    facts = parse_cv_text(
        "Yousef Al-Harbi\nSenior Engineer\nyousef@example.com\n" + EXTRA_SECTIONS)

    assert facts.major_achievements and facts.training_courses and facts.publications
    surgical = next((s for s in facts.additional_sections
                     if "surgic" in s.section_title.lower()), None)
    assert surgical, [s.section_title for s in facts.additional_sections]

    body = " ".join(surgical.entries)
    assert "1,240" in body and "0.8%" in body
    # The same content in a named field AND the catch-all would print twice.
    catch_all = " ".join(e for s in facts.additional_sections for e in s.entries).lower()
    assert "deeplearning.ai" not in catch_all and "journal of applied ai" not in catch_all
